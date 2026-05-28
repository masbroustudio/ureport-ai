import os
import tempfile

import fitz
import pytest
from docx import Document

from app.rag.loaders import load_docx, load_pdf, load_txt


@pytest.fixture
def sample_txt_path():
    return os.path.join(os.path.dirname(__file__), "fixtures", "sample.txt")


@pytest.fixture
def sample_pdf_path():
    """Create a temporary PDF with two pages of text."""
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
        doc = fitz.open()
        # Page 1
        page1 = doc.new_page()
        page1.insert_text((72, 72), "This is the first page of the test PDF document.")
        # Page 2
        page2 = doc.new_page()
        page2.insert_text((72, 72), "This is the second page with additional content.")
        doc.save(f.name)
        doc.close()
        yield f.name
    os.unlink(f.name)


@pytest.fixture
def sample_docx_path():
    """Create a temporary DOCX with multiple paragraphs."""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        doc = Document()
        doc.add_paragraph("First section of content in the document.")
        doc.add_paragraph("More content in the first section.")
        doc.add_paragraph("")  # empty paragraph to create section boundary
        doc.add_paragraph("Second section starts here with new content.")
        doc.add_paragraph("Additional detail in the second section.")
        doc.save(f.name)
        yield f.name
    os.unlink(f.name)


class TestLoadPdf:
    def test_returns_pages(self, sample_pdf_path):
        result = load_pdf(sample_pdf_path)
        assert len(result) == 2

    def test_page_structure(self, sample_pdf_path):
        result = load_pdf(sample_pdf_path)
        for entry in result:
            assert "text" in entry
            assert "page" in entry
            assert "section" in entry
            assert len(entry["text"]) > 0

    def test_page_numbers(self, sample_pdf_path):
        result = load_pdf(sample_pdf_path)
        assert result[0]["page"] == 1
        assert result[1]["page"] == 2

    def test_section_is_none(self, sample_pdf_path):
        result = load_pdf(sample_pdf_path)
        for entry in result:
            assert entry["section"] is None


class TestLoadDocx:
    def test_returns_sections(self, sample_docx_path):
        result = load_docx(sample_docx_path)
        assert len(result) == 2

    def test_section_structure(self, sample_docx_path):
        result = load_docx(sample_docx_path)
        for entry in result:
            assert "text" in entry
            assert "page" in entry
            assert "section" in entry
            assert len(entry["text"]) > 0

    def test_page_is_none(self, sample_docx_path):
        result = load_docx(sample_docx_path)
        for entry in result:
            assert entry["page"] is None

    def test_section_labels(self, sample_docx_path):
        result = load_docx(sample_docx_path)
        assert result[0]["section"] == "section_1"
        assert result[1]["section"] == "section_2"


class TestLoadTxt:
    def test_returns_single_entry(self, sample_txt_path):
        result = load_txt(sample_txt_path)
        assert len(result) == 1

    def test_text_structure(self, sample_txt_path):
        result = load_txt(sample_txt_path)
        entry = result[0]
        assert "text" in entry
        assert "page" in entry
        assert "section" in entry
        assert len(entry["text"]) > 0

    def test_metadata_is_none(self, sample_txt_path):
        result = load_txt(sample_txt_path)
        entry = result[0]
        assert entry["page"] is None
        assert entry["section"] is None

    def test_content_loaded(self, sample_txt_path):
        result = load_txt(sample_txt_path)
        assert "artificial intelligence" in result[0]["text"]
